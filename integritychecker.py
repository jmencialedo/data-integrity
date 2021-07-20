# -*- coding: utf-8 -*-
"""
Created on Thu Jul  8 09:07:42 2021

@author: JavierMencíaLedo
"""

import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
import os
from scipy.stats import pearsonr
from itertools import combinations
from docx import Document
import dataframe_image as dfi
from docx.shared import Inches


def corrmaxmin(df, pairs, maxim):
    if maxim==1:
        result = ["",0]
    else:
        result = ["",1]
    for i in pairs:
        corr, p_value = pearsonr(df[i[0]].tolist(), df[i[1]].tolist())
        if maxim==1 and abs(corr)>abs(result[1]) and i[0]!=i[1]:
            result = [i, corr]
        elif maxim==0 and abs(corr)<abs(result[1]) and i[0]!=i[1]:
            result =  [i,corr]
    return result


def corrcolor(corr):
    color = "#ff0000"
    if abs(corr)>0.9: 
        color = "#00ff00"
    elif abs(corr)>0.75: 
        color ="#58ad2a"
    elif abs(corr)>0.6:
        color ="#ffa600"
    return color

def minuniquepair(headers,unique):
    lowest = [1,""]
    secondl = [1,""]
    for u in range(1,len(unique)):
        if eval(unique[u])<lowest[0]:
            lowest=[eval(unique[u]), headers[u]]
        elif eval(unique[u])<secondl[0]:
            secondl = [eval(unique[u]), headers[u]]
    lp = [lowest[1],secondl[1]]
    return lp

def plotter(df,p, corr, maxmin):
    first = p[0]
    second = p[1]
    x,y = df[first].tolist(),df[second].tolist()
    m, b = np.polyfit(x, y, 1)
    plt.plot(x,y,'x')
    corr, p_value = pearsonr(x, y)
    color = corrcolor(corr)
    if abs(corr)>0.5: plt.plot(x, m*np.array(x)+b, color)
    
    plt.xlabel(first)
    plt.ylabel(second)
    if maxmin==1:
        plt.title('Table of two variables with best correlation: ')
        plt.savefig("bestcorrelation.jpg")
    else:
        plt.title('Table of two variables with worst correlation: ')
        plt.savefig("worstcorrelation.jpg")
    plt.show()
    
def widthchooser(final):
    if len(final.columns)<6:
        width = len(final.columns)
    else:
        width = 6
    return width
    
"""Initial file processing"""
#Three test files:
filepath = "mtcars.csv"
filepath2 = "div.csv"
filepath2 = "superstoreSales.csv"
filepath = "california.csv"

fin = pd.read_csv(filepath)
rpath = "Integrityreport.doc"


"""Data Integrity checks"""
df = pd.DataFrame(fin)
#Remove empty columns
empty_cols = [col for col in df.columns if df[col].isnull().all()]
# Drop these columns from the dataframe
df.drop(empty_cols,
        axis=1,
        inplace=True)
#Column processing    
headers = list(df.columns)
headers.insert(0, "Column Name")
headers[-1]=headers[-1].strip("\n")



dt = ["Data type"] #Creating columns for report file
means, sd,unique, maximum, minimum, totalsum = ["Mean"],["SD"], ["Unique"], ["Maximum"],["Minimum"],["Sum"]
numeric, boolean, booleandata = [],[],[]
modal = list(df.mode().loc[0])
modal.insert(0,"Mode")
mostunique = [headers[1],0]
for h in headers[1:]:
    #headers[headers.index(h)]=h.lower()
    if (len(set(df[h].tolist()))>2) or (len(set(df[h].tolist()))==1):# or len(df[h].tolist())<=2) and h!="id":
        dt.append(df[h].dtype)
    elif h=="id":
        dt.append("Primary Key")
    else:
        print(set(df[h].tolist()))
        dt.append("Boolean")
    if (df[h].dtype == np.int64 or df[h].dtype==np.float64) :
        if df[h].dtype == np.int64: df[h].astype(np.int32)
        else: df[h].astype(np.float32)
        means.append("{:.5}".format(np.nanmean(df[h].tolist())))
        sd.append("{:.4}".format(np.nanstd(df[h].tolist())))
        maximum.append(np.nanmax(df[h].tolist()))
        minimum.append(np.nanmin(df[h].tolist()))
        totalsum.append(df[h].sum())
        if eval(str(len(set(df[h].tolist())))+"/"+str(len(df[h].tolist()))):
            numeric.append(h)
            
    else:
        means.append("N/A")
        sd.append("N/A")
        data = sorted(list(set(sorted([str(i) for i in df[h].tolist()]))))
        if data[-1]=="nan":
            data[:-1]
        maximum.append(data[-1])
        minimum.append(data[0])
        totalsum.append("N/A")
        if eval(str(len(set(df[h].tolist())))+"/"+str(len(df[h].tolist())))>mostunique[1]:
            mostunique = [h, eval(str(len(set(df[h].tolist())))+"/"+str(len(df[h].tolist())))]

    unique.append(str(len(set(df[h].tolist())))+"/"+str(len(df[h].tolist())))
    temp = df[h].dropna()
    temp = [i for i in temp if i is not None]
    if (len(set(temp))==2 and len(temp)>2):
        print(set(temp))
        boolean.append(h)
        booleandata.append(set(temp))
print(mostunique)   
dtypestable = [headers, dt,means, sd, minimum, maximum, unique, modal, totalsum]
final = pd.DataFrame(dtypestable)
"""Generating report"""
doc = Document()
doc.add_heading('Integrity report: '+filepath, 0)
dfi.export(final, 'dataframe.png', max_cols=15)
width = widthchooser(final)
doc.add_picture("dataframe.png", width = Inches(width))


#Check duplicated columns
if len(set(headers))==len(headers):
    doc.add_paragraph("\nRepeated columns? No\n")
else:
    msg = "\nRepeated columns? Yes, "+ ",".join([str(i) for i in list(set([x for x in headers if headers.count(x) > 1]))])
    doc.add_paragraph(msg+"\n")
    
df = df.loc[:,~df.columns.duplicated()]
dup = df[df.duplicated(keep="first")]
doc.add_paragraph("Number of repeated records: "+str(len(dup))+ "\n")
if len(dup)>0: 
    doc.add_paragraph("Repeated records:\n")
    dfi.export(dup, 'duplicates.png', max_rows=11)
    doc.add_picture("duplicates.png", width = Inches(6))   

#Processing duplicates and n
dfcopy = df.copy(deep = True)
column_means = df.interpolate()
for i in df.columns[df.isnull().any(axis=0)]:     #---Applying Only on variables with NaN values
    if i in boolean:
        df[i].fillna(df[i].mode(),inplace=True)
    else:
        df[i].fillna(df[i].interpolate(method='spline', order = 3, limit_direction='forward'),inplace=True)
df = df.drop_duplicates()
dfcopy = dfcopy.drop_duplicates()
df = df.dropna()

"""Calculate correlation between numeric fields"""
pairs = list(combinations(numeric, 2))
if len(pairs)>0:
    p,corr =corrmaxmin(df,pairs,1)
    p2,corr2 = corrmaxmin(df,pairs,0)

    #notunique = minuniquepair(headers, unique)
    """Data visualisation"""
    #Graph with two fields with best corr and table with frequency of column with least uniques
    plotter(df,p, corr,1)#Call to make graph of best correlation
    
    plotter(df,p2, corr2,0)#Call to make graph of worst correlation
    doc.add_picture("bestcorrelation.jpg", width = Inches(6))
    doc.add_paragraph("PMCC: "+str(corr))
    doc.add_picture("worstcorrelation.jpg", width=Inches(6))
    doc.add_paragraph("PMCC: "+str(corr2))
                


"""Report of results"""
file = "results\\"+filepath.split(".")[0]+"(checked).csv"
count = 0
while os.path.isfile(file)==True:
    count+=1
    file = "results\\"+filepath.split(".")[0]+"(" +str(count)+").csv"
df.to_csv(file)
os.startfile(file)


doc.add_paragraph('Any boolean values? ')
if len(boolean)==0:
    doc.add_paragraph("No")
else:
    doc.add_paragraph("Yes: "+", ".join(boolean))
    for b in boolean:
        dfcopy.fillna(0.5, inplace = True)
        tb = pd.DataFrame(dfcopy[b].value_counts())
        before = tb[b].sum()
        while tb.shape[0]>2:
            tb = tb.drop(tb.index[2])
        if  before!= tb[b].sum():
            df2 = pd.DataFrame([before-tb[b].sum()])
            tb.loc[dfcopy.shape[0]] = [before-tb[b].sum()]
            tb.index = ["False","True","nan"]
        else:
            tb.index = ["False", "True"]
        doc.add_paragraph(b, style='Heading 1')
        dfi.export(tb, 'boolean.png')
        doc.add_picture("boolean.png")
title = mostunique[0]

for n in headers[1:]:
    heading = doc.add_paragraph(n.title(), style='Heading 1')
    heading.underline = True
    description = []
    ds = pd.DataFrame(df[n].describe())
    dfi.export(ds, 'description.png')
    doc.add_picture("description.png")
    if n in numeric:
        p = pd.DataFrame(df[n])
        if len(df[n].unique()) !=1:
            minX, maxX = min(df[n].tolist()), max(df[n].tolist())
            p.plot(kind='density', sharey=True, ind=np.linspace(minX, maxX, 700))
            plt.xlim(minX, maxX)
            for i in range(5):
                color = "orange"
                if i==2: color = "red"
                plt.axvline(x=np.quantile(df[n].tolist(),(i)*0.25), color = color)
            plt.title("Distribution of "+n)
            plt.savefig("distribution.jpg")
            plt.show()
            doc.add_picture("distribution.jpg", width = Inches(5))
        x,y = df[title].tolist(), df[n].tolist()
        plt.plot(x,y,'x')

        plt.title(title+" against "+n)
        if len(df[n].tolist())>50:
            plt.gca().xaxis.set_major_locator(plt.NullLocator())
        else:
            plt.xticks(rotation=90)
            plt.tick_params(axis='x', labelsize=8) # Set the x-axis label size
    
        plt.savefig("series.jpg", bbox_inches="tight")
        plt.show()
        doc.add_picture("series.jpg", width = Inches(6))
        
filereport = 'results\\report.docx'
rep = 'results\\report'+filepath.split(".")[0]+'.docx'
count = 0
while os.path.isfile(filereport)==True:
    count+=1
    filereport = rep.split(".")[0]+"(" +str(count)+").docx"
print("Report file name: "+filereport)
doc.save(filereport)
os.startfile(filereport)

