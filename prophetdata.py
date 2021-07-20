# -*- coding: utf-8 -*-
"""
Created on Thu Jul 15 14:07:07 2021

@author: JavierMencíaLedo
"""

import numpy as np
import matplotlib.pyplot as plt 
import pandas as pd
import seaborn as sns
from statsmodels.tsa.api import SimpleExpSmoothing
from fbprophet import Prophet
from fbprophet.plot import add_changepoints_to_plot
from docx import Document
from docx.shared import Inches
import os

def plot_it():
    plt.plot(db2['y'],'go',markersize=2,label='Actual')
    plt.fill_between(
       np.arange(db2.shape[0]), db2['lower'], db2['upper'], alpha=0.5, color="r",
       label="Predicted interval")
    plt.xlabel("Ordered samples.")
    plt.ylabel("Values and prediction intervals.")
    plt.savefig("plusminus.jpg", bbox_inches="tight")
    plt.show()
    doc.add_picture("plusminus.jpg", width = Inches(6))
    
print("Hello world")
filepath = 'NAB/data/realAWSCloudwatch/rds_cpu_utilization_cc0c53.csv'
filepath = 'NAB/data/realAWSCloudwatch/rds_cpu_utilization_e47b3b.csv'
filepath = 'california.csv'
filepath = "bike-sharing-daily.csv"
filepath2 = 'daily-female-births.csv'
filepath2  ="AAPL.csv"
filepath2 = "madrid_2003.csv"
df = pd.read_csv(filepath)[:10000]
print("Read")
rpath = "anomalyreport.doc"
doc = Document()
doc.add_heading('Anomaly report: '+filepath, 0)

#Remove empty columns
empty_cols = [col for col in df.columns if df[col].isnull().all()]
# Drop these columns from the dataframe
df.drop(empty_cols,
        axis=1,
        inplace=True)

headers = list(df.columns)
headers[-1]=headers[-1].strip("\n")
boolean, numeric = [],[]
mostunique = [headers[1],0]
for h in headers:
    temp = df[h].dropna()
    temp = [i for i in temp if i is not None]
    
    if (len(set(temp))==2 and len(temp)>2):
        boolean.append(h)
    elif (df[h].dtype == np.int64 or df[h].dtype==np.float64) :
        numeric.append(h)
    else:
        if eval(str(len(set(df[h].tolist())))+"/"+str(len(df[h].tolist())))>mostunique[1]:
            mostunique = [h, eval(str(len(set(df[h].tolist())))+"/"+str(len(df[h].tolist())))]
df = df.loc[:,~df.columns.duplicated()]
dup = df[df.duplicated(keep="first")]
#Processing duplicates and n
dfcopy = df.copy(deep = True)
column_means = df.interpolate()
for i in df.columns[df.isnull().any(axis=0)]:     #---Applying Only on variables with NaN values
    if i in boolean:
        df[i].fillna(df[i].mode(),inplace=True)
    else:
        df[i].fillna(df[i].interpolate(method='spline', order = 3, limit_direction='forward'),inplace=True)
df = df.drop_duplicates()
dfcopy = dfcopy.drop_duplicates()

df = df.dropna()
xfield = mostunique[0]

for n in numeric:
    doc.add_heading(n)
    db = df.copy(deep=True)
    db = db[["dteday",n]]
    db.columns = ["ds","y"]
    """Simple Moving Average"""
    db2 = db.copy(deep=True)
    db2['SMA'] = db2.iloc[:,1].rolling(window=10).mean()
    db2['diff'] = db2['y'] - db2['SMA']
    db2[['y','SMA']].plot()
    
    plt.savefig("sma.jpg", bbox_inches="tight")
    plt.show()
    db2['diff'].hist()
    plt.title('The distribution of '+n)
    plt.show()
    doc.add_picture("sma.jpg", width=Inches(6))
    """"Exponential Moving Average"""
    db2 = db.copy(deep=True)
    EMAfit = SimpleExpSmoothing(db2['y']).fit(smoothing_level=0.2,optimized=False)
    EMA = EMAfit.forecast(3).rename(r'$\alpha=0.2$')
    db2['EMA'] = EMAfit.predict(start = 10)
    db2['diff'] = db2['y'] - db2['EMA']
    db2[['y','EMA']].plot()
    plt.title("EMA of "+ n)
    plt.savefig("ema.jpg", bbox_inches="tight")
    plt.show()
    doc.add_picture("ema.jpg", width = Inches(6))

    add = -min(db2["EMA"].tolist()) + max(db2["EMA"].tolist())
    db2['upper'] = db2['EMA'] + add//2
    db2['lower'] = db2['EMA'] - add//2
    plot_it()
    # Fitting with default parameters
    
    db_model = Prophet(daily_seasonality=True)
    db_model.fit(db)
    future= db_model.make_future_dataframe(periods=1000, freq='d') #Number of days to predict for
    
    db_model_data=db_model.predict(future)
    forecast = db_model.predict(future)
    fig= db_model.plot(forecast)
    a = add_changepoints_to_plot(fig.gca(), db_model, forecast)
    plt.title("Prophet "+ n)
    plt.savefig("prophet.jpg", bbox_inches="tight")
    plt.show()
    doc.add_picture("prophet.jpg", width = Inches(6))
    
filereport = 'results\\report.docx'
rep = 'results\\anomalyreport'+filepath.split(".")[0]+'.docx'
count = 0
while os.path.isfile(filereport)==True:
    count+=1
    filereport = rep.split(".")[0]+"(" +str(count)+").docx"
print("Report file name: "+filereport)
doc.save(filereport)
os.startfile(filereport)
